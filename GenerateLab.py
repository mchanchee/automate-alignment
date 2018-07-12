"""
Creating .lab (or .txt) files by segmenting a .docx file and according to given .wav files
"""

import os
import docx

#  folder_exists() and file_exists() are re-usable but minor changes must be brought to the error messages


def folder_exists(this_folder: str) -> bool:
    """Check whether this_folder exists"""
    if not os.path.isdir(this_folder):
        print("\n")
        print("{} does not exist.".format(this_folder))
        print("Please make sure you typed the name of the folder containing the wav files correctly and that you are"
              " in the correct directory.\n")
        return False
    else:
        return True


def file_exists(this_file: str) -> bool:
    """Check whether this_file exists"""
    if not os.path.isfile(this_file):
        print("\n")
        print("{} does not exist.".format(this_file))
        print("Please make sure you typed the name of the docx file correctly and that you are in the correct"
              " directory.\n")
        return False
    else:
        return True


#  no_forbidden_characters() and folder_already exists() are re-usable with no changes necessary


def no_forbidden_characters(this_new_item: str) -> bool:
    """Check whether this_new_item (file or folder) contains no forbidden characters"""
    forbidden_characters = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
    folder_name_wrong = any([char in forbidden_characters for char in this_new_item])
    if folder_name_wrong:
        print("\nYou cannot use the following characters in your folder name:\n")
        print(', '.join(forbidden_characters) + '\n')
        return False
    else:
        return True


def folder_already_exists(this_new_folder: str)-> bool:
    """Check whether a folder called this_folder already exists"""
    this_new_folder_already_exists = os.path.isdir(this_new_folder)
    if this_new_folder_already_exists:
        print("\nThe folder '{}' already exists. Please choose another name or delete the existing folder.\n".format(
            this_new_folder))
        return True
    else:
        return False


def generate_lab_files(word_file_name: str, folder_wav_files: str, new_folder_lab_files: str):
    """
    Look at word_file_name and generate the .lab files in new_folder_lab_files
    """

    # Both word_file_name and folder_wav_files exist, there's no forbidden character in new_folder_lab_files and
    # no other folder is alreay called new_folder_lab_files
    if (file_exists(word_file_name) and
            folder_exists(folder_wav_files) and
            no_forbidden_characters(new_folder_lab_files) and
            not folder_already_exists(new_folder_lab_files)):

        # Open the given word document
        doc = docx.Document(word_file_name)

        # Let's consider an empty paragraph contains an empty string or a succession of spaces (here up to 9 spaces)
        # We can make that list bigger if I haven't been able to predict all possible types of meaningless paragraphs
        list_empty_para = [' ' * x for x in range(10)]

        # Extract the text into a list and ignore empty paragraphs (these tend to be at the end of doc)
        text_doc = [doc.paragraphs[i].text for i in range(len(doc.paragraphs))
                    if doc.paragraphs[i].text not in list_empty_para]
        num_para = len(text_doc)

        # Get the names of the wav files
        wav_files = os.listdir(folder_wav_files)
        num_wav = len(wav_files)

        if num_para != num_wav:
            print("\nThere are {} wav files in '{}' but {} paragraphs in '{}'.".format(num_wav, folder_wav_files,
                                                                                       num_para, word_file_name))
            print("The number of wav files must be the same as the number of paragraphs.")
            print("No lab file has been created.")
            print("Please adjust your files/folders accordingly before re-running the code.\n")

        else:
            # Create the folder for the lab files and move into it
            os.mkdir(new_folder_lab_files)
            os.chdir(new_folder_lab_files)
            for i in range(num_wav):
                lab_file_name = wav_files[i][:-4] + ".lab"
                lab_file_content = text_doc[i]
                lab_file = open(lab_file_name, 'w')
                lab_file.write(lab_file_content)
            print("Success!\n")


if __name__ == '__main__':
    print("The paragraphs in the docx file must match the recordings in the folder containing the wav files.")
    word_file_name = input("Enter the name of the docx file: ")
    folder_wav_files = input("Enter the name of the folder containing the wav files: ")
    new_folder_lab_files = input("Enter the name for a new folder to contain the lab files: ")
    generate_lab_files(word_file_name + ".docx", folder_wav_files, new_folder_lab_files)
